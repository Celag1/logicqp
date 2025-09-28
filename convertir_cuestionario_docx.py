#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertir Cuestionario Digital HTML a Word DOCX
Sistema LogicQP - Grupo 6 - Cel@g
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

def add_checkbox_style(doc):
    """Agregar estilo para checkboxes"""
    styles = doc.styles
    checkbox_style = styles.add_style('Checkbox', 1)  # 1 = paragraph style
    checkbox_style.font.name = 'Arial'
    checkbox_style.font.size = Pt(11)

def create_likert_table(doc, title, questions, question_prefix):
    """Crear tabla de escala Likert"""
    doc.add_heading(title, level=3)
    
    # Crear tabla
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pregunta'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'
    
    # Estilo de encabezados
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar preguntas
    for i, question in enumerate(questions, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = question
        row_cells[1].text = '☐'
        row_cells[2].text = '☐'
        row_cells[3].text = '☐'
        row_cells[4].text = '☐'
        row_cells[5].text = '☐'
        
        # Centrar checkboxes
        for j in range(1, 6):
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajustar ancho de columnas
    for row in table.rows:
        row.cells[0].width = Inches(4.0)  # Pregunta
        for i in range(1, 6):
            row.cells[i].width = Inches(0.5)  # Checkboxes

def create_radio_question(doc, question, name, options):
    """Crear pregunta de opción múltiple"""
    doc.add_paragraph(question, style='Heading 4')
    for option in options:
        doc.add_paragraph(f'☐ {option}', style='List Bullet')

def create_text_area(doc, question, name):
    """Crear área de texto para respuestas abiertas"""
    doc.add_paragraph(question, style='Heading 4')
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)

def main():
    """Función principal"""
    print("🚀 Convirtiendo Cuestionario Digital HTML a Word DOCX...")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    add_checkbox_style(doc)
    
    # Título principal
    title = doc.add_heading('📊 CUESTIONARIO DE ENCUESTA LIKERT - SISTEMA LOGICQP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Evaluación de Usabilidad, Eficiencia y Satisfacción del Usuario')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_paragraph('─' * 80)
    
    # Información general
    doc.add_heading('📋 INFORMACIÓN GENERAL', level=1)
    
    info_data = [
        ('Título:', 'Cuestionario de Evaluación del Sistema LogicQP'),
        ('Versión:', '1.0'),
        ('Fecha:', 'Enero 2025'),
        ('Desarrollado por:', 'Grupo 6 - Cel@g'),
        ('Propósito:', 'Evaluar la usabilidad, eficiencia y satisfacción del usuario con el sistema LogicQP')
    ]
    
    for label, value in info_data:
        p = doc.add_paragraph()
        p.add_run(f'{label} ').bold = True
        p.add_run(value)
    
    doc.add_paragraph('─' * 80)
    
    # Instrucciones
    doc.add_heading('🎯 INSTRUCCIONES', level=1)
    
    doc.add_paragraph('Estimado/a Usuario/a:', style='Heading 3')
    doc.add_paragraph('Este cuestionario tiene como objetivo evaluar su experiencia con el sistema LogicQP. Sus respuestas son muy importantes para mejorar el sistema y brindar un mejor servicio.')
    
    doc.add_paragraph('Instrucciones:', style='Heading 4')
    instructions = [
        'Lea cada pregunta cuidadosamente',
        'Marque con una X la opción que mejor represente su experiencia',
        'Use la escala del 1 al 5 donde:',
        '  1 = Muy en desacuerdo / Muy malo',
        '  2 = En desacuerdo / Malo',
        '  3 = Neutral / Regular',
        '  4 = De acuerdo / Bueno',
        '  5 = Muy de acuerdo / Excelente'
    ]
    
    for instruction in instructions:
        doc.add_paragraph(f'• {instruction}', style='List Bullet')
    
    doc.add_paragraph('Tiempo estimado: 15-20 minutos')
    doc.add_paragraph('Confidencialidad: Sus respuestas son completamente anónimas y confidenciales')
    
    doc.add_paragraph('─' * 80)
    
    # Escala de evaluación
    doc.add_heading('📊 ESCALA DE EVALUACIÓN', level=1)
    
    scale_data = [
        ('1 = Muy en desacuerdo / Muy malo', '(0-20%)'),
        ('2 = En desacuerdo / Malo', '(21-40%)'),
        ('3 = Neutral / Regular', '(41-60%)'),
        ('4 = De acuerdo / Bueno', '(61-80%)'),
        ('5 = Muy de acuerdo / Excelente', '(81-100%)')
    ]
    
    for scale, percentage in scale_data:
        p = doc.add_paragraph()
        p.add_run(f'• {scale}').bold = True
        p.add_run(f' {percentage}')
    
    doc.add_paragraph('─' * 80)
    
    # Datos del encuestado
    doc.add_heading('👤 DATOS DEL ENCUESTADO', level=1)
    
    encuestado_data = [
        'Nombre (opcional): _________________________',
        'Rol/Posición: _____________________________',
        'Experiencia con el sistema: _____ meses/años',
        'Fecha de la encuesta: _____________________',
        'Navegador utilizado: _____________________'
    ]
    
    for data in encuestado_data:
        doc.add_paragraph(data)
    
    doc.add_paragraph('Frecuencia de uso:')
    frecuencia = ['Diario', 'Semanal', 'Mensual', 'Ocasional']
    for freq in frecuencia:
        doc.add_paragraph(f'☐ {freq}', style='List Bullet')
    
    doc.add_paragraph('Dispositivo:')
    dispositivos = ['Computadora de escritorio', 'Laptop', 'Tablet', 'Móvil']
    for device in dispositivos:
        doc.add_paragraph(f'☐ {device}', style='List Bullet')
    
    doc.add_paragraph('─' * 80)
    
    # Definir todas las preguntas por sección
    preguntas = {
        'facilidad': [
            'El sistema es fácil de usar',
            'La interfaz es intuitiva y clara',
            'Es fácil navegar por el sistema',
            'Los menús son fáciles de encontrar',
            'Las funciones están bien organizadas'
        ],
        'diseno': [
            'El diseño visual es atractivo',
            'Los colores son apropiados',
            'Los textos son legibles',
            'Los botones son fáciles de identificar',
            'El diseño es consistente en todas las páginas'
        ],
        'navegacion': [
            'Es fácil encontrar lo que busco',
            'La estructura del menú es lógica',
            'Es fácil regresar a páginas anteriores',
            'Los enlaces funcionan correctamente',
            'El sistema tiene un buen flujo de trabajo'
        ],
        'velocidad': [
            'El sistema carga rápidamente',
            'Las páginas se abren sin demoras',
            'Las búsquedas son rápidas',
            'Los reportes se generan rápidamente',
            'El sistema responde bien a mis acciones'
        ],
        'productividad': [
            'Puedo completar mis tareas rápidamente',
            'El sistema me ayuda a ser más productivo',
            'Puedo realizar múltiples tareas simultáneamente',
            'El sistema me ahorra tiempo',
            'Puedo trabajar de manera eficiente'
        ],
        'funcionalidades': [
            'La búsqueda de productos es eficiente',
            'El proceso de compra es rápido',
            'La gestión de inventario es eficaz',
            'Los reportes son útiles y completos',
            'Las notificaciones son oportunas'
        ],
        'satisfaccion': [
            'Estoy satisfecho con el sistema en general',
            'El sistema cumple con mis expectativas',
            'Recomendaría el sistema a otros',
            'El sistema mejora mi experiencia de trabajo',
            'Estoy contento con la calidad del sistema'
        ],
        'experiencia': [
            'El sistema es agradable de usar',
            'Me siento cómodo usando el sistema',
            'El sistema me da confianza',
            'Me siento apoyado por el sistema',
            'El sistema es confiable'
        ],
        'valor': [
            'El sistema aporta valor a mi trabajo',
            'Las funcionalidades son útiles',
            'El sistema resuelve mis necesidades',
            'El sistema es indispensable para mi trabajo',
            'El sistema supera a otros sistemas similares'
        ],
        'productos': [
            'Es fácil buscar productos',
            'Los filtros funcionan bien',
            'La información de productos es clara',
            'Es fácil agregar productos al carrito',
            'La gestión de categorías es eficiente'
        ],
        'compra': [
            'El carrito de compras es fácil de usar',
            'El proceso de checkout es claro',
            'Los formularios son fáciles de llenar',
            'Las opciones de pago son claras',
            'Las confirmaciones son útiles'
        ],
        'inventario': [
            'Es fácil crear nuevos productos',
            'La actualización de stock es sencilla',
            'Las alertas de reposición son útiles',
            'Los reportes de inventario son completos',
            'La trazabilidad de lotes es efectiva'
        ],
        'reportes': [
            'Es fácil generar reportes',
            'Los filtros de reportes son útiles',
            'La exportación de datos funciona bien',
            'Los gráficos son claros y útiles',
            'La información de auditoría es completa'
        ],
        'seguridad': [
            'Me siento seguro usando el sistema',
            'El sistema protege mi información',
            'Los controles de acceso son adecuados',
            'El sistema es seguro para transacciones',
            'Confío en la seguridad del sistema'
        ],
        'confiabilidad': [
            'El sistema funciona de manera consistente',
            'Rara vez experimento errores',
            'El sistema está disponible cuando lo necesito',
            'Los datos se guardan correctamente',
            'El sistema es estable y confiable'
        ],
        'responsivo': [
            'El sistema funciona bien en mi dispositivo',
            'La interfaz se adapta bien a diferentes tamaños',
            'Es fácil usar el sistema en móvil',
            'Los elementos son fáciles de tocar',
            'El sistema es accesible desde cualquier lugar'
        ],
        'accesibilidad': [
            'El sistema es fácil de usar para personas con discapacidades',
            'Los textos tienen buen contraste',
            'El sistema es compatible con lectores de pantalla',
            'Los elementos son fáciles de identificar',
            'El sistema es inclusivo'
        ],
        'problemas': [
            'He experimentado errores en el sistema',
            'El sistema a veces es lento',
            'Algunas funciones no funcionan como esperaba',
            'He tenido problemas de conectividad',
            'El sistema a veces se cuelga'
        ],
        'mejoras': [
            'Necesito más funcionalidades',
            'El sistema necesita mejoras en la interfaz',
            'Necesito mejor rendimiento',
            'El sistema necesita mejor documentación',
            'Necesito mejor soporte técnico'
        ],
        'impacto': [
            'El sistema me ayuda a ser más productivo',
            'El sistema mejora la calidad de mi trabajo',
            'El sistema me ahorra tiempo',
            'El sistema reduce errores en mi trabajo',
            'El sistema mejora la comunicación'
        ],
        'objetivos': [
            'El sistema cumple con los objetivos del negocio',
            'El sistema es rentable para la organización',
            'El sistema mejora la satisfacción del cliente',
            'El sistema reduce costos operativos',
            'El sistema mejora la competitividad'
        ]
    }
    
    # Agregar todas las secciones de preguntas
    secciones = [
        ('🔍 SECCIÓN 1: USABILIDAD', [
            ('1.1 Facilidad de Uso General', 'facilidad'),
            ('1.2 Diseño de la Interfaz', 'diseno'),
            ('1.3 Navegación y Estructura', 'navegacion')
        ]),
        ('⚡ SECCIÓN 2: EFICIENCIA', [
            ('2.1 Velocidad y Rendimiento', 'velocidad'),
            ('2.2 Productividad', 'productividad'),
            ('2.3 Funcionalidades Específicas', 'funcionalidades')
        ]),
        ('😊 SECCIÓN 3: SATISFACCIÓN', [
            ('3.1 Satisfacción General', 'satisfaccion'),
            ('3.2 Experiencia de Usuario', 'experiencia'),
            ('3.3 Valor y Utilidad', 'valor')
        ]),
        ('🔧 SECCIÓN 4: FUNCIONALIDADES ESPECÍFICAS', [
            ('4.1 Gestión de Productos', 'productos'),
            ('4.2 Proceso de Compra', 'compra'),
            ('4.3 Gestión de Inventario', 'inventario'),
            ('4.4 Reportes y Auditoría', 'reportes')
        ]),
        ('🛡️ SECCIÓN 5: SEGURIDAD Y CONFIABILIDAD', [
            ('5.1 Seguridad', 'seguridad'),
            ('5.2 Confiabilidad', 'confiabilidad')
        ]),
        ('📱 SECCIÓN 6: RESPONSIVIDAD Y ACCESIBILIDAD', [
            ('6.1 Diseño Responsivo', 'responsivo'),
            ('6.2 Accesibilidad', 'accesibilidad')
        ]),
        ('🚨 SECCIÓN 7: PROBLEMAS Y MEJORAS', [
            ('7.1 Problemas Encontrados', 'problemas'),
            ('7.2 Necesidades de Mejora', 'mejoras')
        ]),
        ('🎯 SECCIÓN 8: OBJETIVOS DE NEGOCIO', [
            ('8.1 Impacto en el Trabajo', 'impacto'),
            ('8.2 Cumplimiento de Objetivos', 'objetivos')
        ])
    ]
    
    section_counter = 1
    for section_title, subsections in secciones:
        doc.add_heading(section_title, level=1)
        
        for subsection_title, question_key in subsections:
            create_likert_table(doc, subsection_title, preguntas[question_key], question_key)
        
        doc.add_paragraph('─' * 80)
        section_counter += 1
    
    # Agregar sección de comentarios
    doc.add_heading('💬 SECCIÓN 9: COMENTARIOS Y SUGERENCIAS', level=1)
    
    doc.add_paragraph('8.1 Comentarios Libres', style='Heading 3')
    
    comentarios = [
        ('¿Qué es lo que más le gusta del sistema LogicQP?', 'me_gusta'),
        ('¿Qué es lo que menos le gusta del sistema?', 'no_gusta'),
        ('¿Qué funcionalidades le gustaría que se agreguen?', 'funcionalidades_nuevas'),
        ('¿Qué mejoras sugiere para el sistema?', 'mejoras_sugeridas'),
        ('¿Alguna otra observación o comentario?', 'otras_observaciones')
    ]
    
    for pregunta, name in comentarios:
        create_text_area(doc, pregunta, name)
    
    doc.add_paragraph('─' * 80)
    
    # Agregar evaluación general
    doc.add_heading('📊 SECCIÓN 10: EVALUACIÓN GENERAL', level=1)
    
    doc.add_paragraph('9.1 Puntuación General', style='Heading 3')
    doc.add_paragraph('En una escala del 1 al 5, ¿cómo calificaría el sistema LogicQP en general?')
    
    puntuacion_opciones = [
        '1 - Muy malo',
        '2 - Malo',
        '3 - Regular',
        '4 - Bueno',
        '5 - Excelente'
    ]
    for option in puntuacion_opciones:
        doc.add_paragraph(f'☐ {option}', style='List Bullet')
    
    doc.add_paragraph('9.2 Comparación con Otros Sistemas', style='Heading 3')
    doc.add_paragraph('¿Cómo compara LogicQP con otros sistemas similares que ha usado?')
    
    comparacion_opciones = [
        'Mucho peor',
        'Peor',
        'Similar',
        'Mejor',
        'Mucho mejor'
    ]
    for option in comparacion_opciones:
        doc.add_paragraph(f'☐ {option}', style='List Bullet')
    
    doc.add_paragraph('9.3 Recomendación', style='Heading 3')
    doc.add_paragraph('¿Recomendaría LogicQP a otros usuarios?')
    
    recomendacion_opciones = [
        'Definitivamente no',
        'Probablemente no',
        'Neutral',
        'Probablemente sí',
        'Definitivamente sí'
    ]
    for option in recomendacion_opciones:
        doc.add_paragraph(f'☐ {option}', style='List Bullet')
    
    doc.add_paragraph('─' * 80)
    
    # Información adicional
    doc.add_heading('📋 INFORMACIÓN ADICIONAL', level=1)
    
    doc.add_paragraph('Datos del Encuestador', style='Heading 3')
    encuestador_data = [
        'Nombre del encuestador: _________________________',
        'Fecha de administración: _________________________',
        'Hora de inicio: _________________________',
        'Hora de finalización: _________________________',
        'Observaciones: _________________________'
    ]
    
    for data in encuestador_data:
        doc.add_paragraph(data)
    
    doc.add_paragraph('Datos del Sistema', style='Heading 3')
    sistema_data = [
        'Versión del sistema: _________________________',
        'Navegador y versión: _________________________',
        'Sistema operativo: _________________________',
        'Resolución de pantalla: _________________________'
    ]
    
    for data in sistema_data:
        doc.add_paragraph(data)
    
    doc.add_paragraph('─' * 80)
    
    # Instrucciones para el análisis
    doc.add_heading('📊 INSTRUCCIONES PARA EL ANÁLISIS', level=1)
    
    doc.add_paragraph('Escala de Puntuación', style='Heading 3')
    escala_data = [
        '1 = Muy en desacuerdo / Muy malo (0-20%)',
        '2 = En desacuerdo / Malo (21-40%)',
        '3 = Neutral / Regular (41-60%)',
        '4 = De acuerdo / Bueno (61-80%)',
        '5 = Muy de acuerdo / Excelente (81-100%)'
    ]
    
    for data in escala_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_paragraph('Cálculo de Puntuaciones', style='Heading 3')
    calculo_data = [
        'Usabilidad: Promedio de las secciones 1.1, 1.2, 1.3',
        'Eficiencia: Promedio de las secciones 2.1, 2.2, 2.3',
        'Satisfacción: Promedio de las secciones 3.1, 3.2, 3.3',
        'Puntuación Total: Promedio de todas las secciones'
    ]
    
    for data in calculo_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_paragraph('Interpretación de Resultados', style='Heading 3')
    interpretacion_data = [
        '4.5 - 5.0: Excelente (81-100%)',
        '3.5 - 4.4: Bueno (61-80%)',
        '2.5 - 3.4: Regular (41-60%)',
        '1.5 - 2.4: Malo (21-40%)',
        '1.0 - 1.4: Muy malo (0-20%)'
    ]
    
    for data in interpretacion_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_paragraph('─' * 80)
    
    # Notas adicionales
    doc.add_heading('📝 NOTAS ADICIONALES', level=1)
    
    doc.add_paragraph('Espacio para observaciones del encuestador:')
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    
    doc.add_paragraph('Espacio para comentarios del supervisor:')
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    
    doc.add_paragraph('─' * 80)
    
    # Información de contacto
    doc.add_heading('📞 INFORMACIÓN DE CONTACTO', level=1)
    
    doc.add_paragraph('Para consultas sobre esta encuesta:', style='Heading 3')
    contacto_data = [
        'Email: grupo6@celag.com',
        'Teléfono: +593-XX-XXX-XXXX',
        'Sitio web: www.logicqp.com'
    ]
    
    for data in contacto_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_paragraph('Para soporte técnico del sistema:', style='Heading 3')
    soporte_data = [
        'Email: soporte@logicqp.com',
        'Teléfono: +593-XX-XXX-XXXX',
        'Horario: Lunes a Viernes, 8:00 AM - 6:00 PM'
    ]
    
    for data in soporte_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_paragraph('─' * 80)
    
    # Firma y fecha
    doc.add_heading('📄 FIRMA Y FECHA', level=1)
    
    doc.add_paragraph('Firma del encuestado: _________________________')
    doc.add_paragraph('Fecha: _________________________')
    doc.add_paragraph()
    doc.add_paragraph('Firma del encuestador: _________________________')
    doc.add_paragraph('Fecha: _________________________')
    
    doc.add_paragraph('─' * 80)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('© 2025 Grupo 6 - Cel@g. Todos los derechos reservados.')
    doc.add_paragraph('Este cuestionario es parte del proceso de evaluación continua del sistema LogicQP y contribuye al mejoramiento de la experiencia del usuario.')
    
    # Guardar documento
    output_file = 'CUESTIONARIO_DIGITAL_LIKERT_LogicQP.docx'
    doc.save(output_file)
    
    print(f"✅ Cuestionario digital Word creado exitosamente: {output_file}")
    print(f"📊 Total de secciones: 10")
    print(f"📝 Total de preguntas Likert: 100+")
    print(f"⏱️ Tiempo estimado de llenado: 15-20 minutos")
    print(f"📋 Incluye: Usabilidad, Eficiencia, Satisfacción, Funcionalidades, Seguridad, Accesibilidad")
    print(f"📝 Formato: Word DOCX para edición fácil")

if __name__ == "__main__":
    main()


