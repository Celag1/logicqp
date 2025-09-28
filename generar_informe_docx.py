#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import datetime

def add_hyperlink(paragraph, text, url):
    """Agregar un hipervínculo a un párrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Color azul para el enlace
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    # Subrayado
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_informe_docx():
    """Crear el informe en formato Word"""
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos principales
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(102, 126, 234)  # Azul corporativo
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(102, 126, 234)
    subtitle_style.paragraph_format.space_before = Pt(12)
    subtitle_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para código
    code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.right_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Título principal
    title = doc.add_heading('📋 INFORME TÉCNICO: MANEJO DE SESIONES Y COOKIES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Sistema Farmacéutico LogicQP', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('Autor: ').bold = True
    info_para.add_run('Grupo 6 - Cel@g - 2025')
    info_para.add_run('\nFecha: ').bold = True
    info_para.add_run(f'{datetime.datetime.now().strftime("%B %Y")}')
    info_para.add_run('\nVersión: ').bold = True
    info_para.add_run('1.0.0')
    
    # Índice
    doc.add_heading('📑 ÍNDICE', level=1)
    toc_items = [
        'Resumen Ejecutivo',
        'Arquitectura de Autenticación',
        'Gestión de Sesiones',
        'Manejo de Cookies',
        'Almacenamiento Local',
        'Seguridad y Validación',
        'Flujo de Autenticación',
        'Configuración del Sistema',
        'Recomendaciones',
        'Conclusiones'
    ]
    
    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')
    
    # 1. RESUMEN EJECUTIVO
    doc.add_heading('🎯 RESUMEN EJECUTIVO', level=1)
    
    doc.add_paragraph('El sistema LogicQP implementa un sistema robusto de autenticación y gestión de sesiones basado en Supabase Auth, que proporciona:')
    
    # Lista de características
    features = [
        'Autenticación segura con JWT tokens',
        'Sesiones persistentes con auto-refresh',
        'Gestión de perfiles de usuario',
        'Almacenamiento local para datos del carrito',
        'Validación de sesiones en tiempo real',
        'Logout automático por inactividad'
    ]
    
    for feature in features:
        doc.add_paragraph(f'✅ {feature}', style='List Bullet')
    
    doc.add_heading('Tecnologías Utilizadas:', level=2)
    technologies = [
        'Supabase Auth - Autenticación y autorización',
        'Next.js 14 - Framework de aplicación',
        'Zustand - Gestión de estado global',
        'LocalStorage - Persistencia de datos del cliente',
        'JWT Tokens - Tokens de sesión seguros'
    ]
    
    for tech in technologies:
        doc.add_paragraph(f'• {tech}', style='List Bullet')
    
    # 2. ARQUITECTURA DE AUTENTICACIÓN
    doc.add_heading('🏗️ ARQUITECTURA DE AUTENTICACIÓN', level=1)
    
    doc.add_heading('Componentes Principales', level=2)
    
    # Estructura de archivos
    code_para = doc.add_paragraph()
    code_para.add_run('// Estructura de autenticación\n')
    code_para.add_run('├── lib/supabase/\n')
    code_para.add_run('│   ├── client.ts          // Cliente Supabase (Frontend)\n')
    code_para.add_run('│   └── server.ts          // Cliente Supabase (Backend)\n')
    code_para.add_run('├── hooks/\n')
    code_para.add_run('│   ├── useAuth.ts         // Hook principal de autenticación\n')
    code_para.add_run('│   └── use-auth.ts        // Hook alternativo\n')
    code_para.add_run('├── app/api/auth/\n')
    code_para.add_run('│   ├── login/route.ts     // Endpoint de login\n')
    code_para.add_run('│   └── register/route.ts  // Endpoint de registro\n')
    code_para.add_run('└── lib/auth.ts            // Utilidades de autenticación')
    code_para.style = 'CodeStyle'
    
    doc.add_heading('Configuración de Supabase', level=2)
    
    code_para2 = doc.add_paragraph()
    code_para2.add_run('// lib/supabase/client.ts\n')
    code_para2.add_run('export const supabase = createClient(supabaseUrl, supabaseAnonKey, {\n')
    code_para2.add_run('  auth: {\n')
    code_para2.add_run('    autoRefreshToken: true,    // ✅ Auto-renovación de tokens\n')
    code_para2.add_run('    persistSession: true,      // ✅ Persistencia de sesión\n')
    code_para2.add_run('    detectSessionInUrl: true   // ✅ Detección de sesión en URL\n')
    code_para2.add_run('  }\n')
    code_para2.add_run('})')
    code_para2.style = 'CodeStyle'
    
    doc.add_heading('Variables de Entorno', level=2)
    
    code_para3 = doc.add_paragraph()
    code_para3.add_run('# Configuración requerida\n')
    code_para3.add_run('NEXT_PUBLIC_SUPABASE_URL=https://tu-proyecto.supabase.co\n')
    code_para3.add_run('NEXT_PUBLIC_SUPABASE_ANON_KEY=tu_anon_key_aqui\n')
    code_para3.add_run('SUPABASE_SERVICE_ROLE_KEY=tu_service_role_key_aqui\n')
    code_para3.add_run('NEXTAUTH_SECRET=tu_nextauth_secret_key_aqui\n')
    code_para3.add_run('NEXTAUTH_URL=http://localhost:3000')
    code_para3.style = 'CodeStyle'
    
    # 3. GESTIÓN DE SESIONES
    doc.add_heading('🔐 GESTIÓN DE SESIONES', level=1)
    
    doc.add_heading('Hook Principal de Autenticación', level=2)
    
    code_para4 = doc.add_paragraph()
    code_para4.add_run('// hooks/useAuth.ts\n')
    code_para4.add_run('export function useAuth() {\n')
    code_para4.add_run('  const [user, setUser] = useState<User | null>(null)\n')
    code_para4.add_run('  const [session, setSession] = useState<Session | null>(null)\n')
    code_para4.add_run('  const [profile, setProfile] = useState<Profile | null>(null)\n')
    code_para4.add_run('  const [loading, setLoading] = useState(true)\n\n')
    code_para4.add_run('  useEffect(() => {\n')
    code_para4.add_run('    // 1. Obtener sesión inicial\n')
    code_para4.add_run('    supabase.auth.getSession().then(({ data: { session }, error }) => {\n')
    code_para4.add_run('      setSession(session)\n')
    code_para4.add_run('      setUser(session?.user ?? null)\n')
    code_para4.add_run('      if (session?.user) {\n')
    code_para4.add_run('        fetchProfile(session.user.id)\n')
    code_para4.add_run('      }\n')
    code_para4.add_run('    })\n\n')
    code_para4.add_run('    // 2. Escuchar cambios de autenticación\n')
    code_para4.add_run('    const { data: { subscription } } = supabase.auth.onAuthStateChange(\n')
    code_para4.add_run('      async (event, session) => {\n')
    code_para4.add_run('        setSession(session)\n')
    code_para4.add_run('        setUser(session?.user ?? null)\n')
    code_para4.add_run('        if (session?.user) {\n')
    code_para4.add_run('          await fetchProfile(session.user.id)\n')
    code_para4.add_run('        } else {\n')
    code_para4.add_run('          setProfile(null)\n')
    code_para4.add_run('        }\n')
    code_para4.add_run('      }\n')
    code_para4.add_run('    )\n\n')
    code_para4.add_run('    return () => subscription.unsubscribe()\n')
    code_para4.add_run('  }, [])\n')
    code_para4.add_run('}')
    code_para4.style = 'CodeStyle'
    
    doc.add_heading('Estados de Sesión', level=2)
    
    # Tabla de estados
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Estado'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Acción'
    
    # Datos de la tabla
    estados = [
        ('loading: true', 'Verificando sesión', 'Mostrar spinner'),
        ('user: null', 'No autenticado', 'Redirigir a login'),
        ('user: User', 'Autenticado', 'Permitir acceso'),
        ('session: Session', 'Sesión activa', 'Mantener estado')
    ]
    
    for estado, descripcion, accion in estados:
        row_cells = table.add_row().cells
        row_cells[0].text = estado
        row_cells[1].text = descripcion
        row_cells[2].text = accion
    
    # 4. MANEJO DE COOKIES
    doc.add_heading('🍪 MANEJO DE COOKIES', level=1)
    
    doc.add_paragraph('Supabase maneja automáticamente las cookies de sesión:')
    
    code_para5 = doc.add_paragraph()
    code_para5.add_run('// Cookies generadas automáticamente\n')
    code_para5.add_run('const SUPABASE_COOKIES = {\n')
    code_para5.add_run("  'sb-access-token': 'JWT token de acceso',\n")
    code_para5.add_run("  'sb-refresh-token': 'Token de renovación',\n")
    code_para5.add_run("  'sb-provider-token': 'Token del proveedor OAuth'\n")
    code_para5.add_run('}')
    code_para5.style = 'CodeStyle'
    
    doc.add_heading('Persistencia de Sesión', level=2)
    persistencia = [
        'Auto-refresh: Tokens se renuevan automáticamente',
        'Persistencia: Sesión se mantiene entre recargas',
        'Detección: Sesión se detecta en URLs de callback',
        'Limpieza: Cookies se eliminan al cerrar sesión'
    ]
    
    for item in persistencia:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # 5. ALMACENAMIENTO LOCAL
    doc.add_heading('💾 ALMACENAMIENTO LOCAL', level=1)
    
    doc.add_heading('Carrito de Compras (Zustand)', level=2)
    
    code_para6 = doc.add_paragraph()
    code_para6.add_run('// lib/store.ts\n')
    code_para6.add_run('export const useCartStore = create<CartStore>()(\n')
    code_para6.add_run('  persist(\n')
    code_para6.add_run('    (set, get) => ({\n')
    code_para6.add_run('      items: [],\n')
    code_para6.add_run('      total: 0,\n')
    code_para6.add_run('      addItem: (item) => { /* lógica */ },\n')
    code_para6.add_run('      removeItem: (id) => { /* lógica */ },\n')
    code_para6.add_run('      clearCart: () => set({ items: [], total: 0 })\n')
    code_para6.add_run('    }),\n')
    code_para6.add_run("    {\n")
    code_para6.add_run("      name: 'cart-storage', // Clave en localStorage\n")
    code_para6.add_run('    }\n')
    code_para6.add_run('  )\n')
    code_para6.add_run(')')
    code_para6.style = 'CodeStyle'
    
    doc.add_heading('Datos Almacenados Localmente', level=2)
    
    # Tabla de datos almacenados
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Clave'
    hdr_cells2[1].text = 'Tipo'
    hdr_cells2[2].text = 'Descripción'
    
    datos = [
        ('cart-storage', 'Object', 'Carrito de compras'),
        ('user-preferences', 'Object', 'Preferencias del usuario'),
        ('theme-settings', 'String', 'Configuración de tema'),
        ('language-settings', 'String', 'Idioma seleccionado')
    ]
    
    for clave, tipo, descripcion in datos:
        row_cells = table2.add_row().cells
        row_cells[0].text = clave
        row_cells[1].text = tipo
        row_cells[2].text = descripcion
    
    # 6. SEGURIDAD Y VALIDACIÓN
    doc.add_heading('🛡️ SEGURIDAD Y VALIDACIÓN', level=1)
    
    doc.add_heading('Medidas de Seguridad', level=2)
    
    # Tabla de medidas de seguridad
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Medida'
    hdr_cells3[1].text = 'Implementación'
    hdr_cells3[2].text = 'Estado'
    
    medidas = [
        ('JWT Tokens', 'Supabase Auth', '✅ Implementado'),
        ('HTTPS Only', 'Cookies seguras', '✅ En producción'),
        ('Token Refresh', 'Auto-renovación', '✅ Automático'),
        ('Session Timeout', 'Configurable', '✅ 7 días'),
        ('Input Validation', 'Zod schemas', '✅ Implementado'),
        ('Rate Limiting', 'API routes', '⚠️ Pendiente'),
        ('CSRF Protection', 'Next.js built-in', '✅ Automático')
    ]
    
    for medida, implementacion, estado in medidas:
        row_cells = table3.add_row().cells
        row_cells[0].text = medida
        row_cells[1].text = implementacion
        row_cells[2].text = estado
    
    # 7. FLUJO DE AUTENTICACIÓN
    doc.add_heading('🔄 FLUJO DE AUTENTICACIÓN', level=1)
    
    doc.add_heading('Proceso de Login', level=2)
    
    proceso = [
        'Usuario ingresa credenciales',
        'Frontend envía POST a /api/auth/login',
        'API valida credenciales con Supabase',
        'Supabase verifica usuario en base de datos',
        'Se genera JWT + Session',
        'API obtiene perfil del usuario',
        'Se retorna User + Profile + Session',
        'Frontend guarda en estado y redirige'
    ]
    
    for i, paso in enumerate(proceso, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    # 8. CONFIGURACIÓN DEL SISTEMA
    doc.add_heading('⚙️ CONFIGURACIÓN DEL SISTEMA', level=1)
    
    doc.add_heading('Variables de Entorno', level=2)
    
    code_para7 = doc.add_paragraph()
    code_para7.add_run('# .env.local\n')
    code_para7.add_run('NEXT_PUBLIC_SUPABASE_URL=https://tu-proyecto.supabase.co\n')
    code_para7.add_run('NEXT_PUBLIC_SUPABASE_ANON_KEY=eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...\n')
    code_para7.add_run('SUPABASE_SERVICE_ROLE_KEY=eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...\n')
    code_para7.add_run('NEXTAUTH_SECRET=tu_secret_key_muy_seguro\n')
    code_para7.add_run('NEXTAUTH_URL=http://localhost:3000')
    code_para7.style = 'CodeStyle'
    
    # 9. RECOMENDACIONES
    doc.add_heading('🎯 RECOMENDACIONES', level=1)
    
    doc.add_heading('Mejoras de Seguridad', level=2)
    
    mejoras = [
        'Implementar Rate Limiting en API routes',
        'Agregar Middleware de protección de rutas',
        'Implementar autenticación de dos factores (2FA)',
        'Mejorar logs de auditoría de seguridad',
        'Configurar monitoreo de sesiones en tiempo real'
    ]
    
    for mejora in mejoras:
        doc.add_paragraph(f'• {mejora}', style='List Bullet')
    
    doc.add_heading('Optimizaciones de Performance', level=2)
    
    optimizaciones = [
        'Implementar cache de perfiles de usuario',
        'Lazy loading de componentes de autenticación',
        'Sincronización offline de datos',
        'Indicadores de estado de conexión'
    ]
    
    for optimizacion in optimizaciones:
        doc.add_paragraph(f'• {optimizacion}', style='List Bullet')
    
    # 10. CONCLUSIONES
    doc.add_heading('📈 CONCLUSIONES', level=1)
    
    doc.add_paragraph('El sistema LogicQP cuenta con un sistema de autenticación robusto y bien estructurado que cumple con los estándares de seguridad modernos.')
    
    doc.add_heading('Fortalezas', level=2)
    
    fortalezas = [
        'Arquitectura Sólida: Uso de Supabase Auth proporciona una base segura',
        'Persistencia: Sesiones se mantienen entre recargas de página',
        'Auto-refresh: Tokens se renuevan automáticamente',
        'Estado Global: Zustand maneja eficientemente el estado de la aplicación',
        'TypeScript: Tipado fuerte previene errores de runtime'
    ]
    
    for fortaleza in fortalezas:
        doc.add_paragraph(f'• {fortaleza}', style='List Bullet')
    
    doc.add_heading('Áreas de Mejora', level=2)
    
    areas_mejora = [
        'Rate Limiting: Implementar límites de requests por IP',
        'Middleware: Agregar protección de rutas a nivel de servidor',
        '2FA: Implementar autenticación de dos factores',
        'Auditoría: Logs de seguridad más detallados',
        'Monitoreo: Métricas de sesión en tiempo real'
    ]
    
    for area in areas_mejora:
        doc.add_paragraph(f'• {area}', style='List Bullet')
    
    doc.add_heading('Recomendaciones Finales', level=2)
    
    recomendaciones_finales = [
        'Implementar las mejoras de seguridad propuestas en el corto plazo',
        'Configurar monitoreo de sesiones en producción',
        'Realizar auditorías de seguridad periódicas',
        'Documentar procedimientos de recuperación de sesiones',
        'Capacitar al equipo en mejores prácticas de seguridad'
    ]
    
    for recomendacion in recomendaciones_finales:
        doc.add_paragraph(f'• {recomendacion}', style='List Bullet')
    
    # Referencias
    doc.add_heading('📚 REFERENCIAS', level=1)
    
    referencias = [
        'Supabase Auth Documentation - https://supabase.com/docs/guides/auth',
        'Next.js Authentication - https://nextjs.org/docs/authentication',
        'JWT Best Practices - https://tools.ietf.org/html/rfc7519',
        'OWASP Session Management - https://owasp.org/www-community/controls/Session_Management_Cheat_Sheet'
    ]
    
    for ref in referencias:
        doc.add_paragraph(f'• {ref}', style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('© 2025 Grupo 6 - Cel@g. Todos los derechos reservados.')
    
    footer_para = doc.add_paragraph()
    footer_para.add_run('Este informe técnico describe la implementación actual del sistema de autenticación y gestión de sesiones en LogicQP. Para actualizaciones o consultas técnicas, contactar al equipo de desarrollo.').italic = True
    
    # Guardar documento
    doc.save('INFORME_SESIONES_COOKIES_LogicQP.docx')
    print("✅ Documento Word creado exitosamente: INFORME_SESIONES_COOKIES_LogicQP.docx")

if __name__ == "__main__":
    create_informe_docx()


