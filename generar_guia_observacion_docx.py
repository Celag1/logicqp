#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import datetime

def create_guia_observacion_docx():
    """Crear la guía de observación en formato Word"""
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos principales
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(102, 126, 234)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(102, 126, 234)
    subtitle_style.paragraph_format.space_before = Pt(12)
    subtitle_style.paragraph_format.space_after = Pt(6)
    
    # Título principal
    title = doc.add_heading('📋 GUÍA DE OBSERVACIÓN: PROCESOS AS-IS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Sistema Farmacéutico LogicQP', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('Autor: ').bold = True
    info_para.add_run('Grupo 6 - Cel@g - 2025')
    info_para.add_run('\nFecha: ').bold = True
    info_para.add_run(f'{datetime.datetime.now().strftime("%B %Y")}')
    info_para.add_run('\nVersión: ').bold = True
    info_para.add_run('1.0.0')
    info_para.add_run('\nTipo: ').bold = True
    info_para.add_run('Documento de Análisis de Procesos')
    
    # Índice
    doc.add_heading('📑 ÍNDICE', level=1)
    toc_items = [
        'Introducción',
        'Metodología de Observación',
        'Proceso de Ventas AS-IS',
        'Proceso de Inventario AS-IS',
        'Proceso de Auditoría AS-IS',
        'Matriz de Procesos',
        'Hallazgos y Oportunidades',
        'Recomendaciones',
        'Anexos'
    ]
    
    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')
    
    # 1. INTRODUCCIÓN
    doc.add_heading('🎯 INTRODUCCIÓN', level=1)
    
    doc.add_heading('Propósito del Documento', level=2)
    doc.add_paragraph('Esta guía de observación tiene como objetivo documentar el estado actual (AS-IS) de los procesos críticos del sistema LogicQP, específicamente en las áreas de:')
    
    areas = [
        'Ventas y E-commerce',
        'Gestión de Inventario',
        'Auditoría y Control'
    ]
    
    for area in areas:
        doc.add_paragraph(f'• {area}', style='List Bullet')
    
    doc.add_heading('Alcance del Análisis', level=2)
    alcance = [
        'Sistema: LogicQP - Sistema Farmacéutico Inteligente',
        'Módulos: Ventas, Inventario, Auditoría',
        'Usuarios: Administradores, Vendedores, Personal de Inventario',
        'Período: Enero 2025'
    ]
    
    for item in alcance:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('Objetivos Específicos', level=2)
    objetivos = [
        'Mapear los procesos actuales del sistema',
        'Identificar puntos de mejora y optimización',
        'Documentar flujos de trabajo existentes',
        'Establecer baseline para futuras mejoras',
        'Validar cumplimiento de requerimientos'
    ]
    
    for i, objetivo in enumerate(objetivos, 1):
        doc.add_paragraph(f'{i}. {objetivo}', style='List Number')
    
    # 2. METODOLOGÍA DE OBSERVACIÓN
    doc.add_heading('🔍 METODOLOGÍA DE OBSERVACIÓN', level=1)
    
    doc.add_heading('Enfoque Metodológico', level=2)
    
    doc.add_heading('Observación Directa', level=3)
    obs_directa = [
        'Técnica: Shadowing de usuarios',
        'Duración: Sesiones de 2-4 horas por proceso',
        'Frecuencia: 3 sesiones por proceso crítico',
        'Herramientas: Grabación de pantalla, notas detalladas'
    ]
    
    for item in obs_directa:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('Entrevistas Estructuradas', level=3)
    entrevistas = [
        'Participantes: Usuarios finales, administradores',
        'Duración: 45-60 minutos por sesión',
        'Formato: Preguntas abiertas y cerradas',
        'Documentación: Grabación de audio + transcripción'
    ]
    
    for item in entrevistas:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_heading('Análisis de Datos', level=3)
    analisis = [
        'Logs del sistema: Comportamiento de usuarios',
        'Métricas de rendimiento: Tiempos de respuesta',
        'Reportes generados: Calidad y completitud',
        'Errores registrados: Patrones y frecuencia'
    ]
    
    for item in analisis:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    # Tabla de herramientas
    doc.add_heading('Herramientas de Observación', level=2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Herramienta'
    hdr_cells[1].text = 'Propósito'
    hdr_cells[2].text = 'Usuarios'
    
    herramientas = [
        ('Grabación de Pantalla', 'Capturar flujos de trabajo', 'Todos'),
        ('Checklist de Procesos', 'Validar completitud', 'Observadores'),
        ('Formularios de Entrevista', 'Recopilar feedback', 'Usuarios'),
        ('Matriz de Tiempos', 'Medir eficiencia', 'Analistas'),
        ('Mapas de Procesos', 'Visualizar flujos', 'Stakeholders')
    ]
    
    for herramienta, proposito, usuarios in herramientas:
        row_cells = table.add_row().cells
        row_cells[0].text = herramienta
        row_cells[1].text = proposito
        row_cells[2].text = usuarios
    
    # 3. PROCESO DE VENTAS AS-IS
    doc.add_heading('💰 PROCESO DE VENTAS AS-IS', level=1)
    
    doc.add_heading('Descripción General', level=2)
    doc.add_paragraph('El proceso de ventas en LogicQP abarca desde la búsqueda de productos hasta la finalización de la compra, incluyendo gestión del carrito, checkout y confirmación.')
    
    doc.add_heading('Actividades Detalladas', level=2)
    
    doc.add_heading('Búsqueda y Selección de Productos', level=3)
    doc.add_paragraph('Actividad: Navegación del catálogo')
    doc.add_paragraph('Tiempo promedio: 3-5 minutos')
    doc.add_paragraph('Usuarios involucrados: Clientes')
    
    doc.add_paragraph('Pasos observados:', style='Heading 4')
    pasos_busqueda = [
        'Acceso al catálogo (URL: /catalogo)',
        'Carga inicial: 2-3 segundos',
        'Productos mostrados: 12 por página',
        'Aplicación de filtros por categoría, precio, marca',
        'Tiempo de filtrado: 1-2 segundos',
        'Visualización en grid de 4 columnas (desktop)'
    ]
    
    for i, paso in enumerate(pasos_busqueda, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_busqueda = [
        '✅ Positivo: Filtros funcionan correctamente',
        '⚠️ Mejora: Búsqueda por texto no implementada',
        '⚠️ Mejora: No hay comparación de productos'
    ]
    
    for hallazgo in hallazgos_busqueda:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    doc.add_heading('Gestión del Carrito', level=3)
    doc.add_paragraph('Actividad: Agregar/remover productos del carrito')
    doc.add_paragraph('Tiempo promedio: 1-2 minutos')
    doc.add_paragraph('Usuarios involucrados: Clientes')
    
    doc.add_paragraph('Pasos observados:', style='Heading 4')
    pasos_carrito = [
        'Agregar producto con botón "Agregar al carrito"',
        'Confirmación visual: Toast notification',
        'Actualización automática del contador',
        'Modificar cantidades con botones +/-',
        'Input directo de cantidad',
        'Validación: Mínimo 1, máximo stock disponible',
        'Persistencia en LocalStorage',
        'Sincronización en tiempo real'
    ]
    
    for i, paso in enumerate(pasos_carrito, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_carrito = [
        '✅ Positivo: Carrito persistente funciona bien',
        '✅ Positivo: Validación de stock en tiempo real',
        '⚠️ Mejora: No hay guardado de carrito por usuario'
    ]
    
    for hallazgo in hallazgos_carrito:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    doc.add_heading('Proceso de Checkout', level=3)
    doc.add_paragraph('Actividad: Finalización de la compra')
    doc.add_paragraph('Tiempo promedio: 5-8 minutos')
    doc.add_paragraph('Usuarios involucrados: Clientes')
    
    doc.add_paragraph('Pasos observados:', style='Heading 4')
    pasos_checkout = [
        'Revisión del carrito con lista de productos',
        'Cálculo de totales y aplicación de descuentos',
        'Formulario de datos de envío con validación',
        'Campos: Nombre, dirección, teléfono, email',
        'Selección de método de pago',
        'Validación de datos de pago',
        'Resumen de la compra y términos',
        'Confirmación final de la compra'
    ]
    
    for i, paso in enumerate(pasos_checkout, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_checkout = [
        '✅ Positivo: Formulario bien validado',
        '⚠️ Mejora: No hay guardado de direcciones frecuentes',
        '❌ Problema: Pasarela de pago no implementada'
    ]
    
    for hallazgo in hallazgos_checkout:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    # Métricas de rendimiento
    doc.add_heading('Métricas de Rendimiento', level=2)
    
    table_metricas = doc.add_table(rows=1, cols=4)
    table_metricas.style = 'Table Grid'
    
    hdr_cells_metricas = table_metricas.rows[0].cells
    hdr_cells_metricas[0].text = 'Métrica'
    hdr_cells_metricas[1].text = 'Valor Actual'
    hdr_cells_metricas[2].text = 'Objetivo'
    hdr_cells_metricas[3].text = 'Estado'
    
    metricas_ventas = [
        ('Tiempo promedio de compra', '12 minutos', '8 minutos', '⚠️'),
        ('Tasa de abandono de carrito', '35%', '25%', '❌'),
        ('Tiempo de carga del catálogo', '3.2 segundos', '2 segundos', '⚠️'),
        ('Disponibilidad del sistema', '98.5%', '99.5%', '⚠️'),
        ('Satisfacción del cliente', '4.2/5', '4.5/5', '⚠️')
    ]
    
    for metrica, actual, objetivo, estado in metricas_ventas:
        row_cells = table_metricas.add_row().cells
        row_cells[0].text = metrica
        row_cells[1].text = actual
        row_cells[2].text = objetivo
        row_cells[3].text = estado
    
    # 4. PROCESO DE INVENTARIO AS-IS
    doc.add_heading('📦 PROCESO DE INVENTARIO AS-IS', level=1)
    
    doc.add_heading('Descripción General', level=2)
    doc.add_paragraph('El proceso de inventario incluye la gestión de productos, control de stock, alertas de reposición y trazabilidad de lotes farmacéuticos.')
    
    doc.add_heading('Actividades Detalladas', level=2)
    
    doc.add_heading('Gestión de Productos', level=3)
    doc.add_paragraph('Actividad: Registro y actualización de productos')
    doc.add_paragraph('Tiempo promedio: 10-15 minutos por producto')
    doc.add_paragraph('Usuarios involucrados: Administradores, Personal de Inventario')
    
    doc.add_paragraph('Pasos observados:', style='Heading 4')
    pasos_productos = [
        'Creación de producto con formulario completo',
        'Campos: Nombre, descripción, precio, categoría',
        'Subida de imagen del producto',
        'Configuración de stock inicial',
        'Asignación a categoría existente o creación nueva',
        'Configuración de atributos específicos',
        'Configuración de precios base y descuentos'
    ]
    
    for i, paso in enumerate(pasos_productos, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_productos = [
        '✅ Positivo: Formulario completo y validado',
        '✅ Positivo: Gestión de categorías flexible',
        '⚠️ Mejora: No hay importación masiva de productos'
    ]
    
    for hallazgo in hallazgos_productos:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    doc.add_heading('Control de Stock', level=3)
    doc.add_paragraph('Actividad: Monitoreo y actualización de inventario')
    doc.add_paragraph('Tiempo promedio: 5-10 minutos por actualización')
    doc.add_paragraph('Usuarios involucrados: Personal de Inventario')
    
    doc.add_paragraph('Pasos observados:', style='Heading 4')
    pasos_stock = [
        'Actualización manual de stock',
        'Entrada de productos recibidos',
        'Salida por ventas',
        'Ajustes de inventario',
        'Alertas automáticas por stock mínimo',
        'Notificaciones por email/SMS',
        'Dashboard con productos críticos',
        'Registro de número de lote',
        'Fecha de vencimiento',
        'Proveedor y origen'
    ]
    
    for i, paso in enumerate(pasos_stock, 1):
        doc.add_paragraph(f'{i}. {paso}', style='List Number')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_stock = [
        '✅ Positivo: Alertas automáticas funcionan',
        '✅ Positivo: Trazabilidad de lotes implementada',
        '❌ Problema: No hay integración con códigos de barras'
    ]
    
    for hallazgo in hallazgos_stock:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    # 5. PROCESO DE AUDITORÍA AS-IS
    doc.add_heading('🔍 PROCESO DE AUDITORÍA AS-IS', level=1)
    
    doc.add_heading('Descripción General', level=2)
    doc.add_paragraph('El proceso de auditoría incluye el seguimiento de actividades, generación de reportes, control de accesos y cumplimiento de normativas farmacéuticas.')
    
    doc.add_heading('Actividades Detalladas', level=2)
    
    doc.add_heading('Registro de Actividades', level=3)
    doc.add_paragraph('Actividad: Captura de eventos del sistema')
    doc.add_paragraph('Tiempo promedio: Automático')
    doc.add_paragraph('Usuarios involucrados: Sistema, Administradores')
    
    doc.add_paragraph('Eventos registrados:', style='Heading 4')
    eventos = [
        'Autenticación: Login/logout de usuarios',
        'Intentos fallidos de acceso',
        'Cambios de contraseña',
        'Operaciones de datos: Creación, modificación, eliminación',
        'Usuario responsable y timestamp',
        'Transacciones comerciales: Ventas, precios, inventario'
    ]
    
    for evento in eventos:
        doc.add_paragraph(f'• {evento}', style='List Bullet')
    
    doc.add_paragraph('Hallazgos:', style='Heading 4')
    hallazgos_auditoria = [
        '✅ Positivo: Registro automático implementado',
        '✅ Positivo: Información detallada capturada',
        '⚠️ Mejora: No hay retención de logs configurada'
    ]
    
    for hallazgo in hallazgos_auditoria:
        doc.add_paragraph(hallazgo, style='List Bullet')
    
    # 6. MATRIZ DE PROCESOS
    doc.add_heading('📊 MATRIZ DE PROCESOS', level=1)
    
    doc.add_heading('Resumen de Procesos', level=2)
    
    table_procesos = doc.add_table(rows=1, cols=5)
    table_procesos.style = 'Table Grid'
    
    hdr_cells_procesos = table_procesos.rows[0].cells
    hdr_cells_procesos[0].text = 'Proceso'
    hdr_cells_procesos[1].text = 'Complejidad'
    hdr_cells_procesos[2].text = 'Automatización'
    hdr_cells_procesos[3].text = 'Eficiencia'
    hdr_cells_procesos[4].text = 'Prioridad'
    
    procesos = [
        ('Ventas', 'Media', '70%', '75%', 'Alta'),
        ('Inventario', 'Alta', '60%', '80%', 'Alta'),
        ('Auditoría', 'Alta', '50%', '70%', 'Media')
    ]
    
    for proceso, complejidad, automatizacion, eficiencia, prioridad in procesos:
        row_cells = table_procesos.add_row().cells
        row_cells[0].text = proceso
        row_cells[1].text = complejidad
        row_cells[2].text = automatizacion
        row_cells[3].text = eficiencia
        row_cells[4].text = prioridad
    
    # 7. HALLAZGOS Y OPORTUNIDADES
    doc.add_heading('🎯 HALLAZGOS Y OPORTUNIDADES', level=1)
    
    doc.add_heading('Hallazgos Principales', level=2)
    
    doc.add_heading('Fortalezas Identificadas', level=3)
    fortalezas = [
        'Arquitectura sólida: Sistema modular bien diseñado',
        'Funcionalidades core completas: Gestión de productos funcional',
        'Interfaz de usuario intuitiva: Diseño moderno y responsive'
    ]
    
    for fortaleza in fortalezas:
        doc.add_paragraph(f'✅ {fortaleza}', style='List Bullet')
    
    doc.add_heading('Áreas de Mejora', level=3)
    mejoras = [
        'Automatización limitada: Procesos manuales en inventario',
        'Reportes básicos: Análisis limitado de datos',
        'Integración incompleta: Pasarela de pago no implementada'
    ]
    
    for mejora in mejoras:
        doc.add_paragraph(f'⚠️ {mejora}', style='List Bullet')
    
    doc.add_heading('Problemas Críticos', level=3)
    problemas = [
        'Seguridad: No hay autenticación de dos factores',
        'Rendimiento: Tiempos de respuesta lentos',
        'Cumplimiento: No hay validación normativa automática'
    ]
    
    for problema in problemas:
        doc.add_paragraph(f'❌ {problema}', style='List Bullet')
    
    # 8. RECOMENDACIONES
    doc.add_heading('📋 RECOMENDACIONES', level=1)
    
    doc.add_heading('Recomendaciones Prioritarias', level=2)
    
    doc.add_heading('Críticas (Implementar inmediatamente)', level=3)
    criticas = [
        'Implementar autenticación de dos factores',
        'Optimizar rendimiento del sistema',
        'Implementar backup automático de logs'
    ]
    
    for i, critica in enumerate(criticas, 1):
        doc.add_paragraph(f'{i}. {critica}', style='List Number')
    
    doc.add_heading('Importantes (Implementar en 3 meses)', level=3)
    importantes = [
        'Desarrollar motor de búsqueda',
        'Implementar dashboards en tiempo real',
        'Integrar códigos de barras'
    ]
    
    for i, importante in enumerate(importantes, 1):
        doc.add_paragraph(f'{i}. {importante}', style='List Number')
    
    doc.add_heading('Deseables (Implementar en 6 meses)', level=3)
    deseables = [
        'Desarrollar análisis predictivo',
        'Integrar con sistemas externos'
    ]
    
    for i, deseable in enumerate(deseables, 1):
        doc.add_paragraph(f'{i}. {deseable}', style='List Number')
    
    # Plan de implementación
    doc.add_heading('Plan de Implementación', level=2)
    
    fases = [
        ('Fase 1: Estabilización (Mes 1)', [
            'Implementar autenticación 2FA',
            'Optimizar rendimiento',
            'Backup automático de logs'
        ]),
        ('Fase 2: Mejoras de Usuario (Mes 2-3)', [
            'Motor de búsqueda',
            'Dashboards en tiempo real',
            'Mejoras en UX'
        ]),
        ('Fase 3: Automatización (Mes 4-6)', [
            'Códigos de barras',
            'Integración de pagos',
            'Workflows automatizados'
        ]),
        ('Fase 4: Inteligencia (Mes 7-12)', [
            'Análisis predictivo',
            'Machine learning',
            'Integración completa'
        ])
    ]
    
    for fase, tareas in fases:
        doc.add_heading(fase, level=3)
        for tarea in tareas:
            doc.add_paragraph(f'• {tarea}', style='List Bullet')
    
    # 9. ANEXOS
    doc.add_heading('📎 ANEXOS', level=1)
    
    # Anexo A: Checklist de Observación
    doc.add_heading('Anexo A: Checklist de Observación', level=2)
    
    doc.add_heading('Proceso de Ventas', level=3)
    checklist_ventas = [
        'Tiempo de carga del catálogo',
        'Funcionamiento de filtros',
        'Proceso de agregar al carrito',
        'Validación de formularios',
        'Proceso de checkout',
        'Confirmación de compra',
        'Gestión de usuarios',
        'Métodos de pago',
        'Notificaciones al cliente',
        'Manejo de errores'
    ]
    
    for item in checklist_ventas:
        doc.add_paragraph(f'☐ {item}', style='List Bullet')
    
    doc.add_heading('Proceso de Inventario', level=3)
    checklist_inventario = [
        'Creación de productos',
        'Actualización de stock',
        'Gestión de categorías',
        'Alertas de reposición',
        'Trazabilidad de lotes',
        'Reportes de inventario',
        'Control de vencimientos',
        'Gestión de proveedores',
        'Códigos de barras',
        'Importación masiva'
    ]
    
    for item in checklist_inventario:
        doc.add_paragraph(f'☐ {item}', style='List Bullet')
    
    doc.add_heading('Proceso de Auditoría', level=3)
    checklist_auditoria = [
        'Registro de actividades',
        'Generación de reportes',
        'Control de accesos',
        'Gestión de roles',
        'Detección de anomalías',
        'Cumplimiento normativo',
        'Retención de logs',
        'Backup de datos',
        'Seguridad de información',
        'Análisis de patrones'
    ]
    
    for item in checklist_auditoria:
        doc.add_paragraph(f'☐ {item}', style='List Bullet')
    
    # Anexo B: Formulario de Entrevista
    doc.add_heading('Anexo B: Formulario de Entrevista', level=2)
    
    doc.add_heading('Preguntas para Usuarios Finales', level=3)
    doc.add_paragraph('Datos del Entrevistado:')
    doc.add_paragraph('Nombre: _________________________')
    doc.add_paragraph('Rol: ____________________________')
    doc.add_paragraph('Experiencia con el sistema: _____ años/meses')
    doc.add_paragraph('Fecha de entrevista: _____________')
    
    doc.add_paragraph('Preguntas Generales:', style='Heading 4')
    preguntas_generales = [
        '¿Qué tan fácil es usar el sistema LogicQP? (Escala 1-5)',
        '¿Cuáles son las funcionalidades que más utiliza?',
        '¿Qué funcionalidades faltan o necesitan mejora?',
        '¿Cuáles son los principales problemas que enfrenta?',
        '¿Qué mejoras sugeriría para el sistema?',
        '¿El sistema cumple con sus expectativas de trabajo?',
        '¿Recomendaría el sistema a otros usuarios?'
    ]
    
    for i, pregunta in enumerate(preguntas_generales, 1):
        doc.add_paragraph(f'{i}. {pregunta}')
        doc.add_paragraph('Respuesta: _________________________________')
        doc.add_paragraph()
    
    doc.add_paragraph('Preguntas Específicas por Proceso:', style='Heading 4')
    
    doc.add_paragraph('VENTAS:', style='Heading 5')
    preguntas_ventas = [
        '¿Qué tan fácil es encontrar productos en el catálogo?',
        '¿El proceso de agregar al carrito es intuitivo?',
        '¿El checkout es claro y fácil de completar?',
        '¿Ha tenido problemas con el proceso de pago?',
        '¿Las notificaciones son claras y útiles?'
    ]
    
    for i, pregunta in enumerate(preguntas_ventas, 1):
        doc.add_paragraph(f'{i}. {pregunta}')
        doc.add_paragraph('Respuesta: _________________________________')
        doc.add_paragraph()
    
    doc.add_paragraph('INVENTARIO:', style='Heading 5')
    preguntas_inventario = [
        '¿La creación de productos es eficiente?',
        '¿Las alertas de stock son útiles y oportunas?',
        '¿La gestión de categorías es flexible?',
        '¿Los reportes de inventario son completos?',
        '¿Falta alguna funcionalidad importante?'
    ]
    
    for i, pregunta in enumerate(preguntas_inventario, 1):
        doc.add_paragraph(f'{i}. {pregunta}')
        doc.add_paragraph('Respuesta: _________________________________')
        doc.add_paragraph()
    
    doc.add_paragraph('AUDITORÍA:', style='Heading 5')
    preguntas_auditoria = [
        '¿Los reportes son fáciles de generar?',
        '¿La información de auditoría es completa?',
        '¿Los controles de acceso son adecuados?',
        '¿Falta alguna funcionalidad de seguridad?',
        '¿El sistema cumple con normativas?'
    ]
    
    for i, pregunta in enumerate(preguntas_auditoria, 1):
        doc.add_paragraph(f'{i}. {pregunta}')
        doc.add_paragraph('Respuesta: _________________________________')
        doc.add_paragraph()
    
    doc.add_heading('Preguntas para Administradores', level=3)
    doc.add_paragraph('Datos del Entrevistado:')
    doc.add_paragraph('Nombre: _________________________')
    doc.add_paragraph('Cargo: __________________________')
    doc.add_paragraph('Experiencia: _____ años')
    doc.add_paragraph('Fecha de entrevista: _____________')
    
    preguntas_admin = [
        '¿El sistema cumple con los requerimientos del negocio?',
        '¿Qué procesos son más críticos para la operación?',
        '¿Qué métricas son más importantes para el seguimiento?',
        '¿Cuáles son los principales riesgos identificados?',
        '¿Qué mejoras prioritarias recomendaría?',
        '¿El sistema es escalable para el crecimiento?',
        '¿La seguridad del sistema es adecuada?',
        '¿El rendimiento cumple con las expectativas?'
    ]
    
    for i, pregunta in enumerate(preguntas_admin, 1):
        doc.add_paragraph(f'{i}. {pregunta}')
        doc.add_paragraph('Respuesta: _________________________________')
        doc.add_paragraph()
    
    # Anexo C: Instrumento de Observación
    doc.add_heading('Anexo C: Instrumento de Observación', level=2)
    
    doc.add_paragraph('Este instrumento debe ser utilizado durante las sesiones de observación directa para documentar el comportamiento del sistema y los usuarios.')
    
    doc.add_heading('Datos de la Sesión', level=3)
    doc.add_paragraph('Fecha: _________________________')
    doc.add_paragraph('Hora de inicio: _________________')
    doc.add_paragraph('Hora de finalización: ___________')
    doc.add_paragraph('Observador: ____________________')
    doc.add_paragraph('Usuario observado: ______________')
    doc.add_paragraph('Proceso observado: ______________')
    doc.add_paragraph('Navegador utilizado: _____________')
    doc.add_paragraph('Dispositivo: ____________________')
    
    doc.add_heading('Checklist de Observación Detallada', level=3)
    
    doc.add_heading('1. NAVEGACIÓN Y ACCESO', level=4)
    doc.add_paragraph('☐ Tiempo de carga de la página principal')
    doc.add_paragraph('☐ Acceso a diferentes secciones del sistema')
    doc.add_paragraph('☐ Funcionamiento del menú de navegación')
    doc.add_paragraph('☐ Responsive design en diferentes dispositivos')
    doc.add_paragraph('☐ Manejo de errores de navegación')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('2. AUTENTICACIÓN Y SEGURIDAD', level=4)
    doc.add_paragraph('☐ Proceso de login')
    doc.add_paragraph('☐ Validación de credenciales')
    doc.add_paragraph('☐ Gestión de sesiones')
    doc.add_paragraph('☐ Logout y cierre de sesión')
    doc.add_paragraph('☐ Control de accesos por rol')
    doc.add_paragraph('☐ Timeout de sesión')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('3. GESTIÓN DE PRODUCTOS', level=4)
    doc.add_paragraph('☐ Búsqueda de productos')
    doc.add_paragraph('☐ Aplicación de filtros')
    doc.add_paragraph('☐ Visualización de detalles')
    doc.add_paragraph('☐ Agregar al carrito')
    doc.add_paragraph('☐ Modificar cantidades')
    doc.add_paragraph('☐ Persistencia del carrito')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('4. PROCESO DE COMPRA', level=4)
    doc.add_paragraph('☐ Revisión del carrito')
    doc.add_paragraph('☐ Llenado de formularios')
    doc.add_paragraph('☐ Validación de datos')
    doc.add_paragraph('☐ Selección de método de pago')
    doc.add_paragraph('☐ Confirmación de compra')
    doc.add_paragraph('☐ Notificaciones al usuario')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('5. GESTIÓN DE INVENTARIO', level=4)
    doc.add_paragraph('☐ Creación de productos')
    doc.add_paragraph('☐ Actualización de stock')
    doc.add_paragraph('☐ Gestión de categorías')
    doc.add_paragraph('☐ Alertas de reposición')
    doc.add_paragraph('☐ Trazabilidad de lotes')
    doc.add_paragraph('☐ Reportes de inventario')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('6. REPORTES Y AUDITORÍA', level=4)
    doc.add_paragraph('☐ Generación de reportes')
    doc.add_paragraph('☐ Filtros y parámetros')
    doc.add_paragraph('☐ Exportación de datos')
    doc.add_paragraph('☐ Visualización de gráficos')
    doc.add_paragraph('☐ Registro de actividades')
    doc.add_paragraph('☐ Control de accesos')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('7. RENDIMIENTO Y USABILIDAD', level=4)
    doc.add_paragraph('☐ Tiempos de respuesta')
    doc.add_paragraph('☐ Facilidad de uso')
    doc.add_paragraph('☐ Claridad de mensajes')
    doc.add_paragraph('☐ Manejo de errores')
    doc.add_paragraph('☐ Feedback visual')
    doc.add_paragraph('☐ Navegación intuitiva')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('8. PROBLEMAS Y ERRORES', level=4)
    doc.add_paragraph('☐ Errores de validación')
    doc.add_paragraph('☐ Errores de sistema')
    doc.add_paragraph('☐ Problemas de conectividad')
    doc.add_paragraph('☐ Errores de permisos')
    doc.add_paragraph('☐ Problemas de rendimiento')
    doc.add_paragraph('☐ Otros problemas identificados')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('9. SUGERENCIAS Y MEJORAS', level=4)
    doc.add_paragraph('☐ Funcionalidades faltantes')
    doc.add_paragraph('☐ Mejoras en la interfaz')
    doc.add_paragraph('☐ Optimizaciones de proceso')
    doc.add_paragraph('☐ Mejoras de seguridad')
    doc.add_paragraph('☐ Mejoras de rendimiento')
    doc.add_paragraph('☐ Otras sugerencias')
    doc.add_paragraph('Observaciones: _________________________________')
    doc.add_paragraph()
    
    doc.add_heading('10. MÉTRICAS CUANTITATIVAS', level=4)
    doc.add_paragraph('Tiempo total de la sesión: _______ minutos')
    doc.add_paragraph('Número de clics realizados: _______')
    doc.add_paragraph('Número de errores encontrados: _______')
    doc.add_paragraph('Tiempo promedio por tarea: _______ segundos')
    doc.add_paragraph('Número de pasos por proceso: _______')
    doc.add_paragraph('Satisfacción del usuario (1-5): _______')
    doc.add_paragraph()
    
    # Anexo D: Matriz de Evaluación
    doc.add_heading('Anexo D: Matriz de Evaluación', level=2)
    
    doc.add_paragraph('Utilice esta matriz para evaluar cada proceso observado en una escala de 1-5 (1=Muy malo, 5=Excelente)')
    
    table_evaluacion = doc.add_table(rows=1, cols=6)
    table_evaluacion.style = 'Table Grid'
    
    hdr_cells_eval = table_evaluacion.rows[0].cells
    hdr_cells_eval[0].text = 'Criterio'
    hdr_cells_eval[1].text = 'Ventas'
    hdr_cells_eval[2].text = 'Inventario'
    hdr_cells_eval[3].text = 'Auditoría'
    hdr_cells_eval[4].text = 'Promedio'
    hdr_cells_eval[5].text = 'Comentarios'
    
    criterios = [
        ('Facilidad de uso', '', '', '', '', ''),
        ('Eficiencia del proceso', '', '', '', '', ''),
        ('Calidad de la interfaz', '', '', '', '', ''),
        ('Velocidad de respuesta', '', '', '', '', ''),
        ('Manejo de errores', '', '', '', '', ''),
        ('Funcionalidad completa', '', '', '', '', ''),
        ('Seguridad', '', '', '', '', ''),
        ('Escalabilidad', '', '', '', '', ''),
        ('Mantenibilidad', '', '', '', '', ''),
        ('Satisfacción del usuario', '', '', '', '', '')
    ]
    
    for criterio, ventas, inventario, auditoria, promedio, comentarios in criterios:
        row_cells = table_evaluacion.add_row().cells
        row_cells[0].text = criterio
        row_cells[1].text = ventas
        row_cells[2].text = inventario
        row_cells[3].text = auditoria
        row_cells[4].text = promedio
        row_cells[5].text = comentarios
    
    # Anexo E: Glosario de Términos
    doc.add_heading('Anexo E: Glosario de Términos', level=2)
    
    table_glosario = doc.add_table(rows=1, cols=2)
    table_glosario.style = 'Table Grid'
    
    hdr_cells_glosario = table_glosario.rows[0].cells
    hdr_cells_glosario[0].text = 'Término'
    hdr_cells_glosario[1].text = 'Definición'
    
    terminos = [
        ('AS-IS', 'Estado actual del proceso o sistema'),
        ('TO-BE', 'Estado futuro deseado del proceso o sistema'),
        ('KPI', 'Indicador clave de rendimiento (Key Performance Indicator)'),
        ('SLA', 'Acuerdo de nivel de servicio (Service Level Agreement)'),
        ('ROI', 'Retorno de inversión (Return on Investment)'),
        ('UX', 'Experiencia de usuario (User Experience)'),
        ('UI', 'Interfaz de usuario (User Interface)'),
        ('API', 'Interfaz de programación de aplicaciones'),
        ('JWT', 'Token web JSON para autenticación'),
        ('CRUD', 'Crear, Leer, Actualizar, Eliminar (operaciones básicas)'),
        ('FEFO', 'Primero en vencer, primero en salir (First Expired, First Out)'),
        ('RLS', 'Seguridad a nivel de fila (Row Level Security)'),
        ('SPA', 'Aplicación de página única (Single Page Application)'),
        ('PWA', 'Aplicación web progresiva (Progressive Web App)'),
        ('GDPR', 'Reglamento general de protección de datos'),
        ('Farmacovigilancia', 'Monitoreo de efectos adversos de medicamentos'),
        ('Trazabilidad', 'Seguimiento del historial de un producto'),
        ('Lote', 'Conjunto de productos fabricados en las mismas condiciones'),
        ('Stock mínimo', 'Cantidad mínima de inventario antes de reordenar'),
        ('Lead time', 'Tiempo entre la orden y la recepción del producto')
    ]
    
    for termino, definicion in terminos:
        row_cells = table_glosario.add_row().cells
        row_cells[0].text = termino
        row_cells[1].text = definicion
    
    # Anexo F: Referencias
    doc.add_heading('Anexo F: Referencias', level=2)
    
    referencias = [
        'Manual de Usuario LogicQP v1.0',
        'Manual Técnico LogicQP v1.0',
        'Documentación de API LogicQP',
        'Estándares de la industria farmacéutica (FDA, EMA)',
        'Regulaciones locales de salud (MSP Ecuador)',
        'Guías de buenas prácticas de software',
        'Estándares de seguridad ISO 27001',
        'Regulaciones de protección de datos (LOPD)',
        'Normativas de farmacovigilancia',
        'Estándares de trazabilidad farmacéutica'
    ]
    
    for i, referencia in enumerate(referencias, 1):
        doc.add_paragraph(f'{i}. {referencia}', style='List Number')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('© 2025 Grupo 6 - Cel@g. Todos los derechos reservados.')
    
    footer_para = doc.add_paragraph()
    footer_para.add_run('Este documento es confidencial y está destinado únicamente para uso interno del proyecto LogicQP.').italic = True
    
    # Guardar documento
    doc.save('GUIA_OBSERVACION_PROCESOS_ASIS_LogicQP.docx')
    print("✅ Guía de observación Word creada exitosamente: GUIA_OBSERVACION_PROCESOS_ASIS_LogicQP.docx")

if __name__ == "__main__":
    create_guia_observacion_docx()
